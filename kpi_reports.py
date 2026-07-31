import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import os
from datetime import datetime

def render_analytics_dashboard(yearly_data_df):
    st.markdown("### 📈 Phân tích Xu hướng & Thống kê KPI")
    if yearly_data_df.empty:
        st.warning("Chưa có dữ liệu tổng kết năm để vẽ biểu đồ. Vui lòng sang tab 'Tổng kết KPI Cả Năm' và chạy báo cáo trước.")
        return
    
    st.markdown("#### 1. Biểu đồ Tổng kết Phân bổ Xếp Loại (Toàn công ty)")
    # Group by grade
    grade_cols = ['Loại A', 'Loại B', 'Loại C', 'Loại D']
    grade_counts = yearly_data_df[grade_cols].sum().reset_index()
    grade_counts.columns = ['Xếp loại', 'Số lượng']
    
    fig1 = px.pie(grade_counts, names='Xếp loại', values='Số lượng', title="Tỷ lệ Xếp loại các tháng trong Năm", color='Xếp loại', 
                 color_discrete_map={'Loại A':'#2ca02c', 'Loại B':'#1f77b4', 'Loại C':'#ff7f0e', 'Loại D':'#d62728'})
    st.plotly_chart(fig1, use_container_width=True)

    st.markdown("#### 2. Biểu đồ So sánh Điểm Trung bình giữa các Nhân viên")
    # We don't have exact numerical scores in yearly_data_df (only grades). 
    # Let's approximate score: A=95, B=85, C=75, D=60
    def approx_score(row):
        total = row['Loại A']*95 + row['Loại B']*85 + row['Loại C']*75 + row['Loại D']*60
        count = row['Loại A'] + row['Loại B'] + row['Loại C'] + row['Loại D']
        return total / count if count > 0 else 0
    
    yearly_data_df['Điểm TB (Ước tính)'] = yearly_data_df.apply(approx_score, axis=1)
    
    fig2 = px.bar(yearly_data_df, x='Nhân viên', y='Điểm TB (Ước tính)', color='Xếp loại Cả Năm',
                 title="Điểm KPI Trung bình (Ước tính) theo Nhân viên",
                 color_discrete_map={'A (100%)':'#2ca02c', 'B (80%)':'#1f77b4', 'C (60%)':'#ff7f0e'})
    st.plotly_chart(fig2, use_container_width=True)

def generate_department_excel(company_name, month, year, data_rows):
    # data_rows is a list of dicts: STT, HoTen, ChucVu, SoLanTre, SoLanSom, SoLanKhongCC, DiemTruTre, DiemTruSom, DiemTruKhongCC, TongTru, DiemConLai, XepLoai, GhiChu
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{month:02d}.{year}"
    
    # Headers
    ws.merge_cells('A1:D1')
    ws['A1'] = company_name
    ws['A1'].font = Font(bold=True)
    
    ws.merge_cells('A2:M2')
    ws['A2'] = f"BẢNG CHẤM ĐIỂM TÍNH LƯƠNG CBNV THÁNG {month:02d}/{year}"
    ws['A2'].font = Font(bold=True, size=14)
    ws['A2'].alignment = Alignment(horizontal='center')
    
    headers_row1 = ['STT', 'Họ và tên', 'Chức vụ', 'Số lần', '', '', 'Điểm trừ đi trễ', 'Điểm trừ về sớm', 'Điểm trừ không chấm công', 'Tổng điểm trừ', 'Tổng điểm còn lại', 'Xếp loại', 'Ghi chú']
    headers_row2 = ['', '', '', 'Đi trễ', 'Về Sớm', 'Không chấm công', '', '', '', '', '', '', '']
    
    ws.append(headers_row1)
    ws.append(headers_row2)
    
    # Merging headers
    ws.merge_cells('A4:A5')
    ws.merge_cells('B4:B5')
    ws.merge_cells('C4:C5')
    ws.merge_cells('D4:F4')
    ws.merge_cells('G4:G5')
    ws.merge_cells('H4:H5')
    ws.merge_cells('I4:I5')
    ws.merge_cells('J4:J5')
    ws.merge_cells('K4:K5')
    ws.merge_cells('L4:L5')
    ws.merge_cells('M4:M5')
    
    thin = Side(border_style="thin", color="000000")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    
    # Áp dụng font Times New Roman, size 13 cho toàn bộ worksheet
    font_default = Font(name='Times New Roman', size=13)
    font_bold = Font(name='Times New Roman', size=13, bold=True)
    
    ws['A1'].font = Font(name='Times New Roman', size=14, bold=True)
    ws['A2'].font = Font(name='Times New Roman', size=16, bold=True)
    
    for row in ws['A4:M5']:
        for cell in row:
            cell.font = font_bold
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border
            cell.fill = PatternFill(start_color="C6E0B4", end_color="C6E0B4", fill_type="solid") # Light green
            
    # Data
    for i, row in enumerate(data_rows, 1):
        r = [
            i, row.get('HoTen', ''), row.get('ChucVu', ''),
            row.get('SoLanTre', ''), row.get('SoLanSom', ''), row.get('SoLanKhongCC', ''),
            row.get('DiemTruTre', ''), row.get('DiemTruSom', ''), row.get('DiemTruKhongCC', ''),
            row.get('TongTru', ''), row.get('DiemConLai', ''), row.get('XepLoai', ''), row.get('GhiChu', '')
        ]
        ws.append(r)
        for cell in ws[ws.max_row]:
            cell.font = font_default
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
    # Column widths
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 15
    for col in ['D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M']:
        ws.column_dimensions[col].width = 12

    out = BytesIO()
    wb.save(out)
    return out.getvalue()

def set_cell_center(cell):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = 1 # Center

def generate_individual_docx(employee_name, month, year, kpi_score, list_tasks, penalties):
    doc = Document()
    
    # Adjust margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # --- Style definitions ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BẢNG ĐÁNH GIÁ CÔNG VIỆC CỦA CBNV\n")
    r.bold = True
    r.font.size = Pt(14)
    r2 = p.add_run(f"THÁNG {month:02d}/{year}")
    r2.bold = True
    r2.font.size = Pt(14)
    
    # --- Info ---
    doc.add_paragraph(f"Họ và tên: {employee_name}       - Chức danh: .....................      - Ban: .....................")
    
    # Calculate score
    total_penalty_c1 = sum(float(adj.get('DiemDieuChinh', 0)) for adj in penalties if 'chuyên cần' in adj.get('LoaiHanhVi', '').lower() or 'trễ' in adj.get('LoaiHanhVi', '').lower() or 'sớm' in adj.get('LoaiHanhVi', '').lower() or 'công' in adj.get('LoaiHanhVi', '').lower())
    total_penalty_c2 = sum(float(t.get('DiemTru', 0)) for t in list_tasks)
    total_penalty_other = sum(float(adj.get('DiemDieuChinh', 0)) for adj in penalties) - total_penalty_c1
    
    total_penalty_all = abs(total_penalty_c1) + abs(total_penalty_c2) + abs(total_penalty_other)
    final_score = 100 - total_penalty_all
    
    doc.add_paragraph(f"Số điểm trừ cả 2 tiêu chí : {round(total_penalty_all, 1)}            Số điểm hoàn thành cả 2 tiêu chí: {round(final_score, 1)}")
    
    # --- Signatures (using invisible table for perfect alignment) ---
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.autofit = True
    
    c00 = sig_table.cell(0, 0)
    c00.text = "NHÂN VIÊN:__________________"
    c00.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c01 = sig_table.cell(0, 1)
    c01.text = "GIÁM ĐỐC BAN:_________________"
    c01.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c10 = sig_table.cell(1, 0)
    c10.text = "Kiểm tra từ Ban HCNS: ____________"
    c10.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c11 = sig_table.cell(1, 1)
    c11.text = f"Tổng điểm để tính lương: {round(final_score, 1)} % lương"
    c11.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c20 = sig_table.cell(2, 0)
    c20.text = "Phê duyệt của Tổng giám đốc:_______"
    c20.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c21 = sig_table.cell(2, 1)
    c21.text = "Phó Tổng giám đốc phụ trách: _________"
    c21.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p_note = doc.add_paragraph()
    p_note.add_run("*Cơ sở đánh giá kết quả tính lương và xếp loại lao động hàng tháng:\n").bold = True
    p_note.add_run("Mức 1:Từ >91 – 100 điểm : 100% lương, xếp loại A trong tháng.\n")
    p_note.add_run("Mức 2:Từ >81 – 91 điểm : 90% lương, xếp loại B trong tháng.\n")
    p_note.add_run("Mức 3:Từ >71 – 81 điểm : 80% lương, xếp loại C trong tháng.\n")
    p_note.add_run("Mức 4: Dưới 71 điểm : 60% lương, xem xét kỷ luật.")
    
    # --- TABLE 1 ---
    t1 = doc.add_table(rows=2, cols=5)
    t1.style = 'Table Grid'
    t1.autofit = False
    for row in t1.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(1.2)
        row.cells[4].width = Inches(1.2)
    
    cell_0_0 = t1.cell(0, 0)
    cell_0_0.merge(t1.cell(1, 0))
    cell_0_0.text = "TT"
    set_cell_center(cell_0_0)
    
    cell_0_1 = t1.cell(0, 1)
    cell_0_1.merge(t1.cell(1, 1))
    cell_0_1.text = "Tiêu chí 1: Đánh giá việc thực hiện thời gian làm việc"
    set_cell_center(cell_0_1)
    
    cell_0_2 = t1.cell(0, 2)
    cell_0_2.merge(t1.cell(0, 4))
    cell_0_2.text = "Chi tiết từ máy Chấm công"
    set_cell_center(cell_0_2)
    
    t1.cell(1, 2).text = "Đi trễ về sớm (lần)"
    set_cell_center(t1.cell(1, 2))
    t1.cell(1, 3).text = "Quên bấm (lần)"
    set_cell_center(t1.cell(1, 3))
    t1.cell(1, 4).text = "Tổng điểm trừ đtc1"
    set_cell_center(t1.cell(1, 4))
    
    row_cells = t1.add_row().cells
    row_cells[0].text = "1"
    set_cell_center(row_cells[0])
    row_cells[1].text = "Tổng số điểm bị trừ (đtc1) tối đa không quá 15 điểm."
    
    # Count penalties for C1
    tre_som = 0
    quen_bam = 0
    for p in penalties:
        lv = p.get('LoaiHanhVi', '').lower()
        if 'trễ' in lv or 'sớm' in lv: tre_som += 1
        if 'công' in lv: quen_bam += 1
    
    row_cells[2].text = str(tre_som)
    set_cell_center(row_cells[2])
    row_cells[3].text = str(quen_bam)
    set_cell_center(row_cells[3])
    row_cells[4].text = str(abs(total_penalty_c1))
    set_cell_center(row_cells[4])
    
    doc.add_paragraph()
    
    # --- TABLE 2 ---
    t2 = doc.add_table(rows=2, cols=5)
    t2.style = 'Table Grid'
    t2.autofit = False
    for row in t2.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(1.2)
        row.cells[4].width = Inches(1.2)
    
    cell2_0_0 = t2.cell(0, 0)
    cell2_0_0.merge(t2.cell(1, 0))
    cell2_0_0.text = "TT"
    set_cell_center(cell2_0_0)
    
    cell2_0_1 = t2.cell(0, 1)
    cell2_0_1.merge(t2.cell(1, 1))
    cell2_0_1.text = "Tiêu chí 2: Đánh giá mức độ hoàn thành công việc"
    set_cell_center(cell2_0_1)
    
    cell2_0_2 = t2.cell(0, 2)
    cell2_0_2.merge(t2.cell(0, 4))
    cell2_0_2.text = "Điểm trừ nhiệm vụ ko hoàn thành"
    set_cell_center(cell2_0_2)
    
    t2.cell(1, 2).text = "Tgian y/c hoàn thành"
    set_cell_center(t2.cell(1, 2))
    t2.cell(1, 3).text = "Kết quả"
    set_cell_center(t2.cell(1, 3))
    t2.cell(1, 4).text = "Tổng điểm trừ đtc2"
    set_cell_center(t2.cell(1, 4))
    
    for i, t in enumerate(list_tasks, 1):
        r_cells = t2.add_row().cells
        r_cells[0].text = str(i)
        set_cell_center(r_cells[0])
        r_cells[1].text = t.get('TenCV', '')
        r_cells[2].text = str(t.get('TgianYC', ''))
        set_cell_center(r_cells[2])
        r_cells[3].text = str(t.get('KetQua', ''))
        set_cell_center(r_cells[3])
        r_cells[4].text = str(t.get('DiemTru', '0'))
        set_cell_center(r_cells[4])
        
    for i in range(3):
        r_cells = t2.add_row().cells
    
    last_row = t2.add_row().cells
    last_row[0].merge(last_row[2])
    last_row[0].text = "Tổng điểm"
    set_cell_center(last_row[0])
    last_row[3].text = "100"
    set_cell_center(last_row[3])
    last_row[4].text = str(abs(total_penalty_c2))
    set_cell_center(last_row[4])
    
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
def generate_yearly_excel(df, year):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"TongKet_{year}"
    
    # Headers
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
    ws.cell(row=1, column=1, value=f"BẢNG TỔNG KẾT KPI CẢ NĂM {year}")
    
    font_default = Font(name='Times New Roman', size=13)
    font_bold = Font(name='Times New Roman', size=13, bold=True)
    ws.cell(row=1, column=1).font = Font(name='Times New Roman', size=16, bold=True)
    ws.cell(row=1, column=1).alignment = Alignment(horizontal='center', vertical='center')
    
    thin = Side(border_style="thin", color="000000")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    
    # Column headers
    for col_idx, col_name in enumerate(df.columns, 1):
        c = ws.cell(row=3, column=col_idx, value=col_name)
        c.font = font_bold
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border = border
        c.fill = PatternFill(start_color="C6E0B4", end_color="C6E0B4", fill_type="solid")
    
    # Data rows
    for row_idx, row_data in enumerate(df.values, 4):
        for col_idx, val in enumerate(row_data, 1):
            c = ws.cell(row=row_idx, column=col_idx, value=val)
            c.font = font_default
            c.border = border
            c.alignment = Alignment(horizontal='center', vertical='center')
    
    # Adjust column widths
    from openpyxl.utils import get_column_letter
    for col_idx in range(1, len(df.columns) + 1):
        col_letter = get_column_letter(col_idx)
        if col_idx == 1: ws.column_dimensions[col_letter].width = 25
        elif col_idx == 2: ws.column_dimensions[col_letter].width = 25
        else: ws.column_dimensions[col_letter].width = 12
        
    out = BytesIO()
    wb.save(out)
    return out.getvalue()
